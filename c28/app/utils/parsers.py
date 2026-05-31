from pathlib import Path
from typing import Union, List
from abc import ABC, abstractmethod
import re


class BaseParser(ABC):
    @abstractmethod
    def parse(self, file_path: Union[str, Path]) -> str:
        pass


class TXTParser(BaseParser):
    def parse(self, file_path: Union[str, Path]) -> str:
        for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                    return _clean_text(content)
            except UnicodeDecodeError:
                continue
            except Exception:
                break
        return ""


class PDFParser(BaseParser):
    def parse(self, file_path: Union[str, Path]) -> str:
        from pypdf import PdfReader
        
        text_parts = []
        reader = PdfReader(str(file_path))
        
        page_count = len(reader.pages)
        batch_size = 20
        
        for batch_start in range(0, page_count, batch_size):
            batch_end = min(batch_start + batch_size, page_count)
            for page_num in range(batch_start, batch_end):
                page = reader.pages[page_num]
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        cleaned = _clean_text(text)
                        if cleaned:
                            text_parts.append(cleaned)
                except Exception:
                    continue
        
        return "\n\n".join(text_parts)


class WordParser(BaseParser):
    def parse(self, file_path: Union[str, Path]) -> str:
        from docx import Document
        
        try:
            doc = Document(str(file_path))
        except Exception:
            return ""
        
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if text and text.strip():
                cleaned = _clean_text(text)
                if cleaned:
                    text_parts.append(cleaned)
        
        for table in doc.tables:
            try:
                for row in table.rows:
                    row_texts = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        cleaned = _clean_text(" | ".join(row_texts))
                        if cleaned:
                            text_parts.append(cleaned)
            except Exception:
                continue
        
        return "\n\n".join(text_parts)


def _clean_text(text: str) -> str:
    if not text:
        return ""
    
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped:
            cleaned_lines.append(stripped)
    
    return '\n'.join(cleaned_lines)


def get_parser(file_type: str) -> BaseParser:
    parsers = {
        "txt": TXTParser(),
        "pdf": PDFParser(),
        "docx": WordParser(),
        "doc": WordParser(),
    }
    return parsers.get(file_type.lower(), TXTParser())


def parse_document(file_path: Union[str, Path], file_type: str) -> str:
    parser = get_parser(file_type)
    return parser.parse(file_path)
